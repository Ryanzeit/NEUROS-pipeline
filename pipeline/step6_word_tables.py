# =============================================================================
# PUBLICATION 1 — STEP 6: FORMATTED WORD TABLES FOR SUBMISSION
# Run AFTER step5.
# Produces journal-ready Word document with:
#   - Table 1: Patient characteristics by race
#   - Table 1b: Patient characteristics by insurance
#   - Table 2: Fully adjusted regression results (all outcomes)
#   - Table 3: Subgroup analyses by diagnosis category
# Output: Publication1_Tables.docx
# =============================================================================

import subprocess
subprocess.run(["pip", "install", "-q", "pandas", "numpy", "python-docx"], check=False)

from google.colab import drive
drive.mount('/content/drive')

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = '/content/drive/MyDrive/mimic 4/'
OUT  = os.path.join(BASE, 'outputs/')

# -----------------------------------------------------------------------------
# HELPERS
# -----------------------------------------------------------------------------
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, border_size=4):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),  'single')
        border.set(qn('w:sz'),   str(border_size))
        border.set(qn('w:color'),'444444')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def style_header_row(table, row_idx=0, bg='2166AC', font_color='FFFFFF', bold=True):
    row = table.rows[row_idx]
    for cell in row.cells:
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold      = bold
                run.font.color.rgb = RGBColor.from_string(font_color)
                run.font.size = Pt(9)

def style_table(table):
    table.style = 'Table Grid'
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Arial'
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell)
        # Alternate row shading
        if i > 0 and i % 2 == 0:
            for cell in row.cells:
                if not any('race_group' in str(v) or 'Variable' in str(v)
                           for para in cell.paragraphs for run in para.runs for v in [run.text]):
                    set_cell_bg(cell, 'F5F9FF')

def add_table_from_df(doc, df, title, footnote=None, header_bg='2166AC'):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1+len(df), cols=len(df.columns))
    table.style = 'Table Grid'

    # Header
    hdr = table.rows[0]
    for j, col in enumerate(df.columns):
        cell = hdr.cells[j]
        cell.text = str(col)
        set_cell_bg(cell, header_bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True; run.font.color.rgb = RGBColor(255,255,255)
                run.font.size = Pt(9); run.font.name = 'Arial'
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for i, (_, row) in enumerate(df.iterrows()):
        data_row = table.rows[i+1]
        for j, val in enumerate(row):
            cell = data_row.cells[j]
            cell.text = str(val) if not pd.isnull(val) else ''
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(9); run.font.name = 'Arial'
                    # Bold section headers (rows where most values are empty)
                    if j == 0 and str(val) and not str(val).startswith('  '):
                        run.bold = True
            # Alternate shading
            if i % 2 == 0:
                set_cell_bg(cell, 'F5F9FF')
            set_cell_borders(cell)

    # Column widths
    total_width = 9.0
    col_width   = total_width / len(df.columns)
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(col_width)

    if footnote:
        p = doc.add_paragraph()
        run = p.add_run(f"Note: {footnote}")
        run.font.size = Pt(8); run.italic = True

    doc.add_paragraph()

# -----------------------------------------------------------------------------
# BUILD DOCUMENT
# -----------------------------------------------------------------------------
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# Title
title_para = doc.add_heading('Publication 1 — Submission Tables', level=1)
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph(
    '"Racial, Socioeconomic, and Insurance-Based Disparities in Neurosurgical ICU Outcomes: A MIMIC-IV Cohort Analysis"'
)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.italic = True; run.font.size = Pt(11)
doc.add_paragraph()

# -------------------------
# TABLE 1: By Race
# -------------------------
try:
    t1 = pd.read_csv(OUT+'table1_by_race.csv')
    # Rename p_value column
    if 'p_value' in t1.columns:
        t1 = t1.rename(columns={'p_value': 'p-value'})
    add_table_from_df(
        doc, t1,
        title='Table 1. Baseline Characteristics of the Neurosurgical ICU Cohort Stratified by Race/Ethnicity',
        footnote=(
            'Data presented as mean (SD) for continuous variables and N (%) for categorical variables. '
            'P-values from ANOVA for continuous variables and chi-square for categorical variables. '
            'Reference group for regression models: White race, Private insurance. '
            'Elixhauser score calculated using van Walraven weighted method. '
            'GCS = Glasgow Coma Scale; ICU = intensive care unit; LOS = length of stay; SD = standard deviation.'
        )
    )
    print("Table 1 (by race) added.")
except FileNotFoundError:
    print("Table 1 by race not found — run Step 2.")

# -------------------------
# TABLE 1b: By Insurance
# -------------------------
try:
    t1b = pd.read_csv(OUT+'table1_by_insurance.csv')
    if 'p_value' in t1b.columns:
        t1b = t1b.rename(columns={'p_value': 'p-value'})
    add_table_from_df(
        doc, t1b,
        title='Table 1b. Baseline Characteristics Stratified by Insurance Status',
        footnote='See Table 1 footnote for abbreviations and statistical methods.',
        header_bg='4DAC26'
    )
    print("Table 1b (by insurance) added.")
except FileNotFoundError:
    print("Table 1b not found — run Step 2.")

doc.add_page_break()

# -------------------------
# TABLE 2: Regression Results
# -------------------------
try:
    t2 = pd.read_csv(OUT+'table2_regression_results.csv')
    # Add significance stars
    def add_stars(p_str):
        try:
            p = float(p_str)
            if p < 0.001: return f"{p_str} ***"
            elif p < 0.01: return f"{p_str} **"
            elif p < 0.05: return f"{p_str} *"
            else: return p_str
        except:
            return p_str
    if 'p-value' in t2.columns:
        t2['p-value'] = t2['p-value'].apply(add_stars)
    elif 'p_format' in t2.columns:
        t2['p-value'] = t2['p_format'].apply(add_stars)
        t2 = t2.drop(columns=['p_format'], errors='ignore')
    add_table_from_df(
        doc, t2,
        title='Table 2. Adjusted Associations Between Race/Ethnicity, Insurance Status, and Neurosurgical ICU Outcomes',
        footnote=(
            'Models adjusted for age, sex, non-English language, neurosurgical diagnosis category, '
            'calendar year group, Elixhauser comorbidity score, and admission GCS. '
            'aOR = adjusted odds ratio; CI = confidence interval; β = regression coefficient (log-transformed outcome). '
            'Reference groups: White race, Private insurance. '
            '*** p<0.001, ** p<0.01, * p<0.05.'
        ),
        header_bg='7B3294'
    )
    print("Table 2 added.")
except FileNotFoundError:
    print("Table 2 not found — run Step 3c.")

doc.add_page_break()

# -------------------------
# TABLE 3: Subgroup Analysis
# -------------------------
try:
    t3 = pd.read_csv(OUT+'subgroup_mortality_by_diagnosis.csv')
    # Clean and select key columns
    t3['Variable'] = (t3['Variable']
        .str.replace(r"C\(race_group, Treatment\('White'\)\)\[T\.", '', regex=True)
        .str.replace(r"C\(insurance_group, Treatment\('Private'\)\)\[T\.", '', regex=True)
        .str.replace(r"\]",'', regex=True))
    t3['aOR (95% CI)'] = t3.apply(
        lambda r: f"{r['OR']:.2f} ({r['CI_low']:.2f}–{r['CI_high']:.2f})", axis=1)
    t3['p-value'] = t3['p_value'].apply(
        lambda p: ("<0.001 ***" if p<0.001 else
                   f"{p:.3f} **" if p<0.01 else
                   f"{p:.3f} *"  if p<0.05 else f"{p:.3f}"))
    t3_out = t3[['subgroup','Variable','n','aOR (95% CI)','p-value']].copy()
    t3_out.columns = ['Diagnosis Category','Variable','N','aOR (95% CI)','p-value']
    add_table_from_df(
        doc, t3_out,
        title='Table 3. Adjusted Odds Ratios for In-Hospital Mortality by Race/Ethnicity — Stratified by Neurosurgical Diagnosis Category',
        footnote=(
            'Each row represents a separate stratified logistic regression model. '
            'Models adjusted for age, sex, insurance status, non-English language, calendar year, '
            'Elixhauser comorbidity score, and admission GCS. '
            'Subgroups with <100 patients or <10 events excluded. '
            'Reference group: White race. *** p<0.001, ** p<0.01, * p<0.05.'
        ),
        header_bg='D6604D'
    )
    print("Table 3 added.")
except FileNotFoundError:
    print("Table 3 not found — run Step 3b.")

# -------------------------
# TABLE 4: Multiple Comparison Correction
# -------------------------
try:
    t4 = pd.read_csv(OUT+'corrected_pvalues.csv')
    t4['Variable'] = (t4['Variable']
        .str.replace(r"C\(race_group.*?\)\[T\.", '', regex=True)
        .str.replace(r"C\(insurance.*?\)\[T\.", '', regex=True)
        .str.replace(r"\]",'', regex=True))
    t4_out = t4[['outcome','source','Variable','p_value','p_bonferroni','p_fdr_bh','sig_fdr']].copy()
    t4_out.columns = ['Outcome','Model','Variable','Raw p','Bonferroni p','FDR p','Sig (FDR)']
    t4_out['Sig (FDR)'] = t4_out['Sig (FDR)'].map({True:'Yes',False:'No'})
    for col in ['Raw p','Bonferroni p','FDR p']:
        t4_out[col] = t4_out[col].apply(lambda p: "<0.001" if float(p)<0.001 else f"{float(p):.3f}" if pd.notna(p) else '')
    add_table_from_df(
        doc, t4_out,
        title='Table 4. Multiple Comparison Correction — Bonferroni and FDR-Adjusted p-Values',
        footnote=(
            'FDR = false discovery rate (Benjamini-Hochberg method). '
            'Bonferroni correction applied across all race/ethnicity and insurance comparisons simultaneously. '
            'Significance threshold: p<0.05 after correction.'
        ),
        header_bg='E08214'
    )
    print("Table 4 added.")
except FileNotFoundError:
    print("Table 4 not found — run Step 3b.")

# -------------------------
# SAVE
# -------------------------
out_path = OUT + 'Publication1_Tables.docx'
doc.save(out_path)
print(f"\n=== STEP 6 COMPLETE ===")
print(f"Word document saved: {out_path}")
print("Contains: Table 1 (by race), Table 1b (by insurance), Table 2 (regression), Table 3 (subgroup), Table 4 (MCC)")
